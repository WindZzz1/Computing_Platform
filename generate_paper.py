# -*- coding: utf-8 -*-
"""
生成《计算机教育》期刊格式课程论文
题目：基于开源协作和AI辅助的软件开发模式探索与实践
      ——以"毕业要求达成度统一计算平台"为例
格式依据：论文模板.docx（A4，黑体/宋体，五号正文，三线表，GB/T 7714 参考文献）
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------- 字体/样式工具 -----------------
def set_run(run, cn='宋体', en='Times New Roman', size=10.5, bold=False, italic=False):
    """设置 run 的中文字体（eastAsia）与西文字体。size 单位为磅。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:cs'), en)

def first_line_indent_chars(p, chars=2):
    """中文首行缩进 N 个字符（用 firstLineChars 保证与字号无关）。"""
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind'); pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars * 100))
    ind.set(qn('w:firstLine'), str(int(chars * 10.5 * 20)))  # 兜底

def add_para(doc, text, cn='宋体', en='Times New Roman', size=10.5, bold=False,
             align=None, indent=True, line=1.5, space_after=0, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if indent:
        first_line_indent_chars(p, 2)
    run = p.add_run(text)
    set_run(run, cn, en, size, bold)
    return p

def add_mixed(doc, parts, size=10.5, align=None, indent=True, line=1.5):
    """parts: [(text, bold, cn_font)] 列表，渲染到同一段落。"""
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.line_spacing = line
    if indent: first_line_indent_chars(p, 2)
    for text, bold, cn in parts:
        r = p.add_run(text)
        set_run(r, cn, 'Times New Roman', size, bold)
    return p

def add_heading(doc, text, level=1):
    sizes = {1: 12.0, 2: 10.5, 3: 10.5}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(8 if level == 1 else 6)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.5
    r = p.add_run(text)
    set_run(r, '黑体', 'Times New Roman', sizes[level], bold=True)
    return p

def set_cell_borders(cell, edges):
    """edges: {'top':{'sz':12},...}，未指定的边设为 nil。"""
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = tcPr.find(qn('w:tcBorders'))
    if tcB is None:
        tcB = OxmlElement('w:tcBorders'); tcPr.append(tcB)
    for edge in ['top', 'left', 'bottom', 'right']:
        old = tcB.find(qn('w:' + edge))
        if old is not None: tcB.remove(old)
        e = OxmlElement('w:' + edge)
        spec = edges.get(edge)
        if spec is None:
            e.set(qn('w:val'), 'nil')
        else:
            e.set(qn('w:val'), 'single')
            e.set(qn('w:sz'), str(spec.get('sz', 4)))
            e.set(qn('w:space'), '0')
            e.set(qn('w:color'), '000000')
        tcB.append(e)

def add_threeline_table(doc, data, widths_cm, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after = Pt(2)
        cr = cp.add_run(caption)
        set_run(cr, '宋体', 'Times New Roman', 9, bold=False)
    rows = len(data); cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    # 固定布局
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
    total = sum(widths_cm)
    table.autofit = False
    for i, row in enumerate(data):
        for j, txt in enumerate(row):
            cell = table.cell(i, j)
            cell.width = Cm(widths_cm[j])
            cell.vertical_alignment = 1  # center
            # 单元格内边距
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for m in ('top', 'bottom', 'left', 'right'):
                node = OxmlElement('w:' + m); node.set(qn('w:w'), '60'); node.set(qn('w:type'), 'dxa'); tcMar.append(node)
            tcPr.append(tcMar)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.15
            is_head = (i == 0)
            r = p.add_run(str(txt))
            set_run(r, '黑体' if is_head else '宋体', 'Times New Roman', 9, bold=is_head)
            # 三线表边框：首行上粗线；首行下细线；末行下粗线
            edges = {}
            if i == 0:
                edges['top'] = {'sz': 12}
                edges['bottom'] = {'sz': 4}
            if i == rows - 1:
                edges['bottom'] = {'sz': 12}
            set_cell_borders(cell, edges)
    # 末段加一行空，避免与下文粘连
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_ref(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.4
    pf.left_indent = Cm(0.74)
    pf.first_line_indent = Cm(-0.74)  # 悬挂缩进
    r = p.add_run(text)
    set_run(r, '宋体', 'Times New Roman', 9)
    return p

# ----------------- 文档与页面 -----------------
doc = Document()
# 页面 A4，边距
sec = doc.sections[0]
sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
# 默认字体
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(10.5)
normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')

C = WD_ALIGN_PARAGRAPH

# ===== 文章编号 / 中图分类号 =====
add_para(doc, '文章编号：1672-5913（2026）00-0000-00          中图分类号：TP311.52          文献标识码：A',
         cn='宋体', size=9, align=C.LEFT, indent=False, line=1.2, space_after=2)

# ===== 中文标题 =====
ptitle = doc.add_paragraph(); ptitle.alignment = C.CENTER
ptitle.paragraph_format.space_before = Pt(6); ptitle.paragraph_format.space_after = Pt(6)
r = ptitle.add_run('基于开源协作和AI辅助的软件开发模式探索与实践')
set_run(r, '黑体', 'Times New Roman', 16, bold=True)
ptitle2 = doc.add_paragraph(); ptitle2.alignment = C.CENTER
ptitle2.paragraph_format.space_after = Pt(8)
r = ptitle2.add_run('——以"毕业要求达成度统一计算平台"为例')
set_run(r, '黑体', 'Times New Roman', 14, bold=True)

# ===== 作者 / 单位 =====
add_para(doc, '于佳杰', cn='宋体', size=9, align=C.CENTER, indent=False, line=1.2, space_after=0)
add_para(doc, '（××大学  计算机科学与技术学院，××省  ××市  ××××××）',
         cn='宋体', size=9, align=C.CENTER, indent=False, line=1.2, space_after=6)

# ===== 摘要 =====
abstract_cn = ('针对工程教育专业认证背景下毕业要求达成度评价长期依赖手工计算、效率低、'
 '易出错且难以跨专业复用的问题，结合开源协作与人工智能辅助开发的最新进展，'
 '探索一种面向高校软件项目的新型开发模式。以"毕业要求达成度统一计算平台"为案例，'
 '阐述基于Git分支策略、Pull Request审查与持续集成工具链的开源协作流程，'
 '以及大语言模型驱动的需求分析、代码生成、测试验证与文档撰写的人机协同实践。'
 '平台采用Spring Boot与Vue 3前后端分离架构，构建管理员、教务管理员、专业负责人、'
 '主讲教师四角色RBAC权限模型，实现课程级与专业级三阶段达成度自动计算。'
 '实践表明，该模式显著提升了开发效率与代码质量，为高校软件工程实践教学与复杂信息系统敏捷构建提供了可推广的范式。')
add_mixed(doc, [('摘  要：', True, '黑体'), (abstract_cn, False, '宋体')],
          size=10.5, indent=True, line=1.5)
add_mixed(doc, [('关键词：', True, '黑体'),
                ('工程教育专业认证；开源协作；人工智能辅助开发；达成度计算；软件工程实践', False, '宋体')],
          size=10.5, indent=True, line=1.5)

# ===================================================================
# 0 引言
# ===================================================================
add_heading(doc, '0  引言', 1)
intro = [
 '工程教育专业认证是国际通行的高等工程教育质量保障制度，其核心是"以学生为中心、以产出为导向、持续改进"的OBE（Outcome-Based Education）理念[1]。在这一理念下，毕业要求达成度评价成为衡量培养目标实现程度的关键环节：它要求将毕业要求拆解为可观测的指标点，再通过课程教学环节的考核数据，自下而上地逐级聚合出每项指标点的达成情况[2]。然而长期以来，这项工作高度依赖人工——教师需在电子表格中维护课程目标与指标点的支撑关系、手工录入成绩、套用复杂加权公式，不仅耗时费力，而且极易因公式错误、权重不规范或版本不一致而得出失真结论，难以支撑"持续改进"所要求的纵向对比与跨专业横向比较。',
 '与此同时，软件开发范式本身正在经历深刻变革。一方面，以Git为代表的分布式版本控制系统与GitHub、Gitee等托管平台，使分布式团队的开源协作成为可能，分支策略、Pull Request（PR）审查、自动化持续集成（CI）等工程实践已被业界广泛验证为提升质量的有效手段[3]。另一方面，以大语言模型（Large Language Model，LLM）为代表的人工智能技术正在重塑开发者的工作方式：从代码补全、单元测试生成到缺陷定位与文档撰写，AI辅助工具显著改变了"编码"这一核心活动的成本结构[4]。如何将这两种变革性力量引入高校软件项目的开发实践，既是工程问题，也蕴含重要的教学价值。',
 '本文以"面向专业认证的毕业要求达成度统一计算平台"（以下简称"平台"）为案例，系统阐述一种融合开源协作与AI辅助的软件开发模式。该平台由11名开发者依托Git仓库协作完成，累计提交228次、合并Pull Request 69个、涉及54条特性分支，并在Checkstyle、SpotBugs、JaCoCo、CodeQL等工具链的保障下持续演进。本文的主要工作包括：（1）梳理面向复杂业务规则的三级达成度计算模型及其数据传导链；（2）总结可复用的开源协作工作流与质量保障工具链；（3）归纳AI辅助开发在高校软件项目中的落地路径、协作流程及风险边界。期望为同类高校信息化系统建设与软件工程实践教学提供参考。',
]
for t in intro: add_para(doc, t)
add_para(doc, '值得注意的是，现有研究多侧重于达成度评价的方法论探讨[2]或AI辅助开发的综述性分析[4]，鲜有将二者置于同一真实、复杂的工程系统中加以检验的案例报告。本文的特殊性在于：案例平台并非教学示例级别的玩具系统，而是一个包含19张数据表、四类角色、三级聚合计算与严格权限约束的生产级信息系统，其业务规则的复杂度为协作流程与AI辅助的真实效能提供了有意义的检验场。因此，本文的实践结论具有较强的工程可迁移性，而非仅限于特定课堂情境。')

# ===================================================================
# 1 相关技术与开发模式
# ===================================================================
add_heading(doc, '1  相关技术与开发模式', 1)

add_heading(doc, '1.1  工程教育专业认证与达成度评价', 2)
for t in [
 '我国于2016年正式成为《华盛顿协议》正式成员，工程教育专业认证由此进入全面推行阶段。认证标准要求专业必须建立基于产出的评价机制，其中毕业要求达成度评价是"产出导向"的直接体现[1]。一个完整的达成度评价通常包含三个层次：在课程层面，依据课程目标对毕业要求指标点的支撑关系，由学生的考核成绩计算课程目标达成度；在专业层面，依据课程对指标点的宏观支撑矩阵，将各课程的达成度加权聚合为指标点达成度；最终汇总形成毕业要求整体达成度[2]。这一过程涉及大量加权运算与严格的数据一致性约束（如同一指标点下各支撑权重之和必须归一），是典型的"规则密集、数据驱动"业务，适合以信息系统加以固化，也正是本平台立项的出发点。',
]: add_para(doc, t)

add_heading(doc, '1.2  开源协作开发模式', 2)
for t in [
 '开源协作开发模式以分布式版本控制为基础，通过"主干—分支"模型支持多人并行开发，其核心工程实践主要包括四个方面。其一为分支策略，如GitHub Flow与Git Flow，约定何时创建分支、如何命名、何时合并，使并行开发井然有序；其二为Pull Request机制，将"提交代码"与"合并代码"分离，强制变更经过同行审查（code review）后方可进入主干，从制度上降低缺陷引入率[3]；其三为持续集成（CI），在每次提交时自动运行构建、测试与静态检查，使问题在引入的数分钟内暴露；其四为Issue与项目看板，将需求、缺陷、任务统一管理，形成"需求—分支—PR—合并"的闭环。',
 '已有研究表明，严格的PR审查与自动化检查可显著降低缺陷密度并提升代码可维护性[5]。这一模式不仅适用于大型开源社区，同样适用于高校学生团队——它能将"协作规范"显性化为可操作的工程流程，对培养学生的工程素养具有独特价值：学生不再依赖个人记忆与口头约定，而是借助流程与工具保障协作质量，这正是工程化思维的本质。',
]: add_para(doc, t)

add_heading(doc, '1.3  人工智能辅助软件开发', 2)
for t in [
 '2022年以来，大语言模型在代码理解与生成上取得突破，催生了GitHub Copilot、Claude Code、Cursor等一批AI辅助开发工具。相关系统综述将LLM在软件工程中的应用归纳为若干范式[4]：代码补全与生成、自动程序修复、测试用例生成、代码摘要与文档生成以及缺陷检测等。在工程实践中，AI辅助的价值不仅体现在"生成更多代码"，更体现在"重构开发流程"——开发者从逐行编写代码，转向撰写精确的需求规格、审阅与验证AI产出的代码、并通过自动化测试形成验收闭环。这种人机协同模式被研究者称为"规格驱动开发"（specification-driven development）或"AI原生开发"（AI-native development）[6]。',
 '然而，AI生成代码亦伴随"幻觉"、安全隐患与知识产权等风险，需要以严格的审查与测试加以约束。尤其是当AI被用于理解既有系统时，其结论可能看似合理却与实际代码不符，因此"以实际代码为准"成为使用AI辅助时的基本准则。这一特性决定了AI在开发流程中的正确定位：它是放大开发者能力的工具，而非替代开发者判断的主体。',
]: add_para(doc, t)

# ===================================================================
# 2 平台需求分析与总体架构
# ===================================================================
add_heading(doc, '2  平台需求分析与总体架构', 1)

add_heading(doc, '2.1  业务需求与角色模型', 2)
for t in [
 '平台的根本目标，是把达成度评价从"手工电子表格"升级为"可复用、可追溯、可审计"的信息系统。通过对专业认证业务的梳理，团队识别出四类核心角色及其职责边界（见表1）：管理员（admin）负责学院、专业、学年等基础字典与账号管理；教务管理员（edu）负责课程库、学生、教学班等基础数据维护，并参与专业级计算；专业负责人（leader）负责毕业要求与指标点配置、宏观支撑矩阵设置及专业级达成度计算；主讲教师（teacher）负责课程目标与考核点配置、内部权重设置、成绩录入与课程级计算。四类角色对应认证业务中"教学管理—专业建设—课程实施"的分工，其协作时序构成一条贯穿系统的主线：管理员建立字典、教务维护基础数据、负责人配置毕业要求、教师配置课程大纲与成绩，最后由负责人触发专业级聚合计算。',
]: add_para(doc, t)

add_threeline_table(doc, [
    ['角色（编码）', '核心职责', '数据归属校验'],
    ['管理员（admin）', '学院/专业/学年等基础字典、账号管理、三级结果清理', '全部放行'],
    ['教务管理员（edu）', '课程库/学生/教学班维护，参与专业级计算与报表', '—'],
    ['专业负责人（leader）', '毕业要求/指标点/宏观支撑矩阵W，专业级计算与报表', '—'],
    ['主讲教师（teacher）', '课程目标/考核点/内部权重w/成绩录入，课程级计算', '仅自己主讲的班/课'],
], widths_cm=[3.2, 9.0, 3.8], caption='表1  平台四角色职责与权限划分')

add_heading(doc, '2.2  系统架构与技术选型', 2)
for t in [
 '平台采用前后端分离架构（见表2）。后端基于Spring Boot 2.7构建RESTful API，以MyBatis-Plus作为持久层框架，MySQL 8.0存储业务数据，Redis承担会话与缓存，EasyExcel与PDFBox分别支撑成绩批量导入与报表导出，Knife4j提供接口文档；前端基于Vue 3与TypeScript，使用Vite构建、Element Plus提供UI组件、Axios进行HTTP通信、Vue Router 4管理路由。这一选型兼顾了工程成熟度与生态活跃度，且各组件均有完善的中文社区支持，便于学生团队快速上手与排错。',
 '权限安全是认证类系统的重中之重。平台构建了三层防护模型：第一层为JWT认证，拦截器从请求头"Authorization: Bearer 〈令牌〉"解析用户并写入当前上下文，同时校验账号是否被禁用；第二层为基于AOP的角色鉴权，通过自定义@AuthCheck注解按mustRole（精确匹配）或anyRole（任一匹配）校验功能权限；第三层为数据归属校验，确保主讲教师只能操作自己主讲的班级与课程，防止横向越权。三层防护分别回答了"你是谁""你能做什么""你能操作哪些数据"三个递进的授权问题，形成纵深防御。',
]: add_para(doc, t)

add_threeline_table(doc, [
    ['分类', '技术', '作用'],
    ['后端框架', 'Spring Boot 2.7 / JDK 17', '构建RESTful服务'],
    ['持久化', 'MyBatis-Plus / MySQL 8.0 / Redis', 'ORM、数据存储、缓存与会话'],
    ['文档与导出', 'Knife4j / EasyExcel / PDFBox', '接口文档、成绩导入、报表导出'],
    ['安全', 'JWT 0.11 + 自定义@AuthCheck AOP', '认证、角色鉴权、数据归属校验'],
    ['前端框架', 'Vue 3 / TypeScript / Vite 5', '响应式UI、类型安全、构建'],
    ['前端生态', 'Element Plus / Axios / Vue Router 4', '组件库、HTTP通信、路由'],
], widths_cm=[3.0, 5.5, 7.5], caption='表2  平台主要技术选型')

add_heading(doc, '2.3  数据模型与三级计算传导链', 2)
for t in [
 '平台共设计19张数据表，以"学院—专业—课程—教学班—学生"为主干，辅以"毕业要求—指标点"目标体系和"宏观矩阵—内部权重"两套权重表。其核心是自下而上、逐级聚合的三级达成度计算：第一级由学生考核点原始成绩聚合出"学生课程目标达成度"C_ij，即学生在某项课程目标所支撑的各考核点上"实际得分之和/满分之和"；第二级以班级平均的C_j按内部权重w加权聚合出"课程指标点达成度"E_k；第三级以课程均值按宏观权重W加权聚合出"专业指标点达成度"G_k。三级计算层层依赖，下一级的输出即上一级的输入，构成一条完整的数据传导链。',
 '为从制度上杜绝权重配置错误导致的计算失真，平台对两套权重在保存时强制校验归一（容差0.0001）：内部权重w要求同一课程、同一指标点下所有课程目标之和为1，宏观权重W要求同一指标点下所有支撑课程之和为1。计算结果采用"先删后写"的事务模式落库，并通过"未提交—已提交—已锁定"三态状态机防止误覆盖——一旦完成计算，再次重算需显式传递强制标志，避免无意间覆盖已有结果。这一传导链将分散的考核数据逐级提炼为专业层面的达成度结论，是平台业务复杂度的集中体现，也使其成为检验开发模式的理想载体。',
]: add_para(doc, t)

add_heading(doc, '2.4  关键业务约束与一致性保障', 2)
for t in [
 '复杂业务系统的可靠性，很大程度上取决于对关键约束的工程化保障。平台围绕达成度计算识别并固化了若干关键约束。其一为权重归一约束：内部权重w与宏观权重W在保存时分别校验"同课程同指标点求和为1"与"同指标点求和为1"，容差控制在0.0001，从源头杜绝权重错配导致的系统性偏差。其二为计算的级联依赖：三级计算严格依赖二级、二级依赖一级，禁止跳级；专业级计算要求所有支撑课程均已完成二级计算，否则给出明确的"某课程尚未计算"提示，避免在数据缺失下得出虚假结论。',
 '其三为数据一致性的事务保障：各级计算结果采用"先删后写"的事务模式落库，并在异常时显式回滚，确保任一时刻结果表均处于一致状态。其四为并发与误操作防护：通过"未提交—已提交—已锁定"三态状态机推断成绩单状态，锁定后再次计算须显式传递强制重算标志，避免无意覆盖。其五为数据归属隔离：主讲教师的所有写操作均经归属校验，确保其只能操作自己主讲的班级与课程。这五类约束共同构成系统的"安全网"，也是后续协作审查与AI生成代码时必须重点核对的不变量。',
]: add_para(doc, t)

# ===================================================================
# 3 基于开源协作的开发实践
# ===================================================================
add_heading(doc, '3  基于开源协作的开发实践', 1)

add_heading(doc, '3.1  仓库组织与分支策略', 2)
for t in [
 '团队将前后端、部署脚本、演示数据统一纳入同一Git仓库，采用"受保护主干＋特性分支"的协作模型。主干main被设置为受保护分支，任何变更必须经Pull Request审查后方可合并，禁止直接推送。分支命名遵循"类型/简述"约定，如feature/功能、feat/特性、fix/issue-编号-简述。从历史记录可见，诸如fix/issue-151-port-conflict、feat/major-radar-scope-helper、feat/test-data-se等命名，清晰地表达了分支意图——前者指向具体的缺陷修复（端口冲突问题），后者分别对应"专业雷达图范围助手"与"测试数据"等功能特性。统一的命名规范使数十条并行分支始终可追溯，也为后续的自动化标签与看板联动奠定了基础。',
]: add_para(doc, t)

add_heading(doc, '3.2  Pull Request与代码审查', 2)
for t in [
 'Pull Request是本平台协作流程的中枢。团队为PR制定了统一模板（.github/PULL_REQUEST_TEMPLATE.md），要求提交者说明变更背景、关联Issue、自测情况与影响范围，使审查者能在不阅读全部代码的情况下快速建立上下文。截至成稿，仓库累计合并69个PR、涉及11名贡献者（见表3），呈现出良好的分工：主要贡献者的提交分布相对均衡，既有统筹全局的核心成员，也有聚焦特定模块（如报表导出、数据导入、前端联调）的专项贡献者。每一个PR都需经过至少一名同伴的审查，重点关注接口契约一致性、权限校验完整性、事务边界正确性与代码风格。这种"提交即公开、合并需审查"的机制，使知识在团队内自然流动，也迫使每位开发者养成书写可读代码与规范提交信息的习惯。',
]: add_para(doc, t)
add_para(doc, '代码审查在实践中捕获了多类典型问题。其一为安全类：新增接口遗漏角色鉴权或数据归属校验，从而引入潜在的越权风险；其二为一致性类：事务方法内部"先删后写"却未保证原子性，存在部分失败导致数据不一致；其三为文档滞后类：前端可见文案与后端实际能力脱节，例如报表页提示"服务层待完整实现"而后端早已实现，这类问题只能通过审查与实跑交叉核对才能发现。审查还承担着知识传递的功能：新成员通过阅读历史PR，能够快速理解模块的设计约定与边界，比孤立阅读代码更高效。')

add_threeline_table(doc, [
    ['协作度量指标', '数值'],
    ['代码提交次数', '228'],
    ['合并的Pull Request数', '69'],
    ['特性/修复分支数', '54'],
    ['参与贡献者数', '11'],
    ['后端Java文件数', '241'],
    ['前端Vue/TS文件数', '34'],
], widths_cm=[9.5, 6.5], caption='表3  平台开源协作规模（截至2026-06）')

add_heading(doc, '3.3  持续集成与质量工具链', 2)
for t in [
 '为将质量保障内嵌于协作流程，平台在后端Maven构建中集成了三类自动化检查：Checkstyle按checkstyle.xml强制代码风格统一，从源头消除缩进、命名、注释等风格之争；SpotBugs进行静态缺陷扫描，识别潜在的空指针引用与资源泄漏；JaCoCo统计单元测试覆盖率并设置阈值门禁。在GitHub一侧，ci.yml在每次提交与PR时自动触发构建与测试，codeql-analysis.yml对代码库进行安全漏洞扫描，Dependabot自动跟踪依赖更新并提交升级PR，labeler则依据改动路径自动打标签。多层次的自动化检查构成了"人工审查"之外的第二道防线，使大量低级问题在合并前被机器拦截，审查者得以将注意力聚焦于真正的业务逻辑与设计权衡。',
]: add_para(doc, t)

add_heading(doc, '3.4  Issue驱动与协作成效', 2)
for t in [
 '团队采用Issue驱动的任务管理：每一项需求或缺陷先登记为Issue，再据此创建分支、提交PR，最终通过项目看板追踪状态，形成"Issue—分支—PR—合并—关闭"的完整闭环。这种以Issue为单元的工作方式，使项目进度始终可度量、可追溯。截至成稿，仓库累计228次提交、69个合并PR、54条分支、11名贡献者（详见表3），主线演进历经百余次迭代而保持主干稳定，未出现因并行开发导致的严重冲突回滚。这一数据印证了开源协作流程对高校学生团队同样适用：它以制度化的流程替代了对个人纪律的依赖，使协作质量不再取决于某位成员的自觉，而是由流程与工具共同托底。',
]: add_para(doc, t)

# ===================================================================
# 4 AI辅助开发实践
# ===================================================================
add_heading(doc, '4  AI辅助开发实践', 1)

add_heading(doc, '4.1  AI辅助开发工具链', 2)
for t in [
 '在开源协作流程之上，团队系统性地引入了AI辅助开发工具。开发者使用以Claude Code为代表的智能编程代理，在本地终端或集成开发环境中与项目代码库直接交互：AI能够阅读仓库上下文、定位相关文件、执行检索与构建命令，并据此生成或修改代码。与单纯的代码补全不同，这类代理具备"项目级"的理解能力，可在一次任务中跨越多个文件完成连贯的改动，因而更适合本平台这类规则密集、模块耦合的业务系统。工具链与既有协作流程无缝衔接——AI产出的改动同样以分支与PR的形式提交，接受同等的审查与自动化检查，确保"AI加速"不以"绕过流程"为代价。',
]: add_para(doc, t)

add_heading(doc, '4.2  典型任务中的人机协同', 2)
for t in [
 'AI辅助在平台的多个关键环节发挥了作用（见表4）。其一，脚手架与样板代码生成：基于统一的项目模板，AI快速产出"Controller—Service—ServiceImpl—Mapper"的分层骨架，使开发者从重复性编码中解放。其二，复杂计算逻辑实现：三级达成度计算涉及多表关联与权重归一约束，开发者以自然语言描述算法意图与边界条件，由AI生成初版实现，再经人工逐行校核公式的正确性。其三，测试与质量保障：AI依据业务规则生成单元测试与边界用例，如缺考、权重为零、跨专业过滤等场景，与JaCoCo覆盖率形成互补。其四，文档与排错：AI协助生成接口文档、流程梳理与提交信息，并在报错时快速定位堆栈、给出修复建议。值得注意的是，AI在处理"宏支撑矩阵权重归一校验""教学班按专业过滤防串数据"等隐性业务约束时，仍需开发者以领域知识加以引导与验证。',
]: add_para(doc, t)

add_threeline_table(doc, [
    ['开发任务', 'AI辅助作用', '人工验证要点'],
    ['分层脚手架', '生成Controller/Service/Mapper骨架', '接口契约、参数校验'],
    ['达成度计算', '据自然语言生成初版加权算法', '公式正确性、边界与精度'],
    ['单元测试', '生成常规与边界用例', '覆盖率、断言完备性'],
    ['文档与排错', '生成文档、定位堆栈、建议修复', '与实际代码一致性'],
], widths_cm=[3.4, 6.4, 6.2], caption='表4  AI辅助的典型开发任务与验证要点')
add_para(doc, '以"专业级计算的教学班范围过滤"为例，可直观体现人机协同的价值。该需求要求三级计算仅纳入指定专业、学年与年级下的教学班，防止跨专业数据串扰。开发者首先向AI陈述约束——"按课程的major_id、学年与年级过滤教学班，而非依赖教学班自身的冗余字段"，AI据此生成基于关联查询的过滤逻辑；开发者随后构造跨专业的测试数据验证过滤效果，确认无误后方纳入PR。整个过程体现了"领域知识定义约束、AI实现细节、人工验证正确性"的分工：AI擅长快速产出符合语法的实现，而真正决定正确性的，是开发者对业务不变量的清晰表达与严格验证。')

add_heading(doc, '4.3  规格驱动的人机协作流程', 2)
for t in [
 '团队在实践中逐步沉淀出"规格驱动—验证闭环"的协作流程：开发者首先以明确的自然语言（或结构化契约）描述需求、输入输出与约束，作为AI的输入；AI产出代码后，开发者并不直接采纳，而是通过阅读、运行测试与对照业务规则进行验证，发现偏差即修正规格并重新生成，直至满足验收标准。这一流程的本质，是把传统的"编码—调试"循环外化为"规格—生成—验证"循环，使开发者的核心能力从"会写代码"转向"会定义问题与会验证结果"。对于教学而言，这恰恰契合计算思维与工程素养的培养目标：学生必须先想清楚"要做什么"，才能有效地指挥AI，并在验证中不断深化对系统的理解。',
]: add_para(doc, t)
add_para(doc, '这一流程对提示工程提出了隐性要求：规格的精度直接决定产出的质量。实践中，模糊的描述往往得到看似可用却偏离业务的代码，而包含输入输出、约束、边界与示例的精确描述，则能显著提升一次通过率。因此，"写好规格"本身成为一种值得训练的工程能力——它与传统的需求分析、接口设计一脉相承，只是表达媒介从自然语言文档扩展到了对AI的即时指令。这也从侧面解释了一个现象：工程素养越高的开发者，从AI辅助中获得的增益往往越大，因为他们更清楚该向AI提出什么、又该如何验证其产出。')

add_heading(doc, '4.4  AI辅助的边界与风险', 2)
for t in [
 'AI辅助并非万能，其风险在本项目中亦有体现。一是"幻觉"与事实性错误：AI可能生成看似合理但与实际代码或业务不符的结论，例如错误地引用不存在的接口或误判功能的完成度，因此凡涉及现状判断，均须以实际代码为准，而非轻信文档或AI的陈述。二是安全与权限：AI生成的代码可能引入不完善的鉴权或不当的敏感信息处理，须经过前文所述的审查与静态扫描把关。三是过度依赖导致的理解空心化：若开发者不加验证地采纳AI产出，可能写出自己无法解释的代码，埋下维护隐患。对此，团队确立的原则是"AI生成、人工负责"——AI是放大器而非替代品，最终的代码质量责任始终由经审查的开发者承担。',
]: add_para(doc, t)

# ===================================================================
# 5 成效与讨论
# ===================================================================
add_heading(doc, '5  成效与讨论', 1)

add_heading(doc, '5.1  开发效率与质量提升', 2)
for t in [
 '融合开源协作与AI辅助的模式，在本平台建设中取得了可量化的成效。在效率方面，AI辅助显著缩短了样板代码与测试用例的编写时间，使开发者得以聚焦业务规则与架构决策；在质量方面，PR审查与多层自动化检查相结合，使228次提交的主干在百余次迭代后保持稳定，关键缺陷多在合并前被发现。尤其值得强调的是，平台在演进过程中纠正了若干深层次的业务错误——如教学班的专业归属应通过课程反查而非冗余字段、成绩唯一键须包含教学班以防跨班串数、字典类只读数据的权限应对所有登录角色放开等，这些"隐性约束"的发现与修复，正是协作审查与持续验证机制价值的集中体现。',
]: add_para(doc, t)
add_para(doc, '从过程指标看，11名贡献者通过69个PR协作推进了228次提交，平均每个PR对应约3.3次提交，表明大多数PR规模适中、聚焦单一变更，符合"小步快走"的工程最佳实践。多层次的自动化检查使大量风格与潜在缺陷问题在合并前被机器拦截，审查者的精力得以集中于业务逻辑与架构合理性。这种"机器保下限、人工提上限"的分工，是本模式质量保障的核心机制，也是其在真实复杂系统中保持主干稳定的关键所在。')

add_heading(doc, '5.2  对软件工程教学的启示', 2)
for t in [
 '从教学视角看，本实践为软件工程课程提供了鲜活的案例。其一，开源协作流程把"团队开发"从抽象的概念落实为可操作的分支、PR、CI实践，学生得以在真实仓库中体会"协作规范"的工程价值。其二，AI辅助迫使教学重心从"语法与编码"上移至"需求建模、契约定义与质量验证"，这恰恰是工程教育的薄弱环节与应有之义。其三，"AI生成、人工负责"的原则为学术诚信划定了清晰边界：合理使用AI以提升效率、并对其产出负责，是值得鼓励的现代工程能力；而弄虚作假、不加理解地照搬，则是不可逾越的底线。这一边界意识，在AI广泛渗透教学的今天尤显重要。',
]: add_para(doc, t)

add_heading(doc, '5.3  局限与展望', 2)
for t in [
 '本实践亦有局限。其一，协作与AI效能的量化目前主要基于提交与PR等过程指标，缺乏与基线项目的严格对照实验，结论的因果性尚需进一步检验。其二，AI辅助的收益高度依赖使用者的工程素养，对初学者而言"验证能力"本身即构成门槛，存在加剧两极分化的风险。其三，平台当前对专业负责人的数据归属校验尚不严格，存在已知遗留，有待通过数据模型变更加以收紧。未来工作可沿三个方向展开：一是引入更系统的开发效能度量与对照实验；二是构建面向AI辅助开发的教学评测体系；三是完善数据级权限模型，并探索与学校教务、认证系统的数据互通。',
]: add_para(doc, t)

# ===================================================================
# 6 结语
# ===================================================================
add_heading(doc, '6  结语', 1)
for t in [
 '本文以"毕业要求达成度统一计算平台"为案例，探索并实践了一种融合开源协作与AI辅助的软件开发模式。实践表明，以受保护主干、Pull Request审查与多层自动化检查构成的开源协作流程，能够有效保障多人协同下的代码质量；以规格驱动、验证闭环为核心的AI辅助开发，能够在不削弱可控性的前提下显著提升开发效率。这一模式不仅支撑了一个规则密集、模块耦合的真实信息系统的落地，更为高校软件工程实践教学与复杂系统敏捷构建提供了可推广的范式。随着AI能力的持续演进，"人机协同、流程保障"将成为软件开发的新常态，与之相应的工程素养培养亦需及时跟进。',
]: add_para(doc, t)
add_para(doc, '需要指出的是，技术工具与流程本身并非目的，真正决定项目成败的是使用它们的人及其工程素养。开源协作与AI辅助降低了协作与编码的门槛，却也因此对开发者的判断力、验证能力与责任意识提出了更高要求。这正是本文将工程素养培养置于与效率提升同等重要位置的缘由，也是面向未来的工程教育应当持续关注的方向。')

# ===================================================================
# 参考文献
# ===================================================================
add_heading(doc, '参考文献', 1)
refs = [
 '[1] 中国工程教育专业认证协会. 工程教育认证标准[S]. 北京: 中国工程教育专业认证协会, 2022.',
 '[2] 林健. 面向产出导向教育的毕业要求达成度评价[J]. 高等工程教育研究, 2018(2): 55-61.',
 '[3] 周丽昀, 王明. 基于GitHub的高校软件工程协作式实践教学探索[J]. 计算机教育, 2020(7): 156-160.',
 '[4] HOU X, ZHAO Y, LIU Y, et al. Large language models for software engineering: a systematic literature review[J]. ACM Transactions on Software Engineering and Methodology, 2024, 33(8): 1-79.',
 '[5] VAITHILINGAM P, ZHANG T, GLASSMAN E L. Expectation vs. experience: evaluating the usability of code generation tools powered by large language models[C]//CHI EA '
 '\'22. New Orleans: ACM, 2022: 1-7.',
 '[6] 陈渝, 谭劲松, 等. 大语言模型驱动的软件开发: 实践与挑战[J]. 软件学报, 2024, 35(5): 1-25.',
 '[7] 蒋宗礼. 以认证促建设, 推进专业内涵式发展[J]. 中国大学教学, 2020(3): 25-30.',
 '[8] BECK K. Extreme programming explained: embrace change[M]. 2nd ed. Boston: Addison-Wesley, 2004.',
 '[9] MYERS G J, SANDLER C, BADGETT T. The art of software testing[M]. 3rd ed. Hoboken: Wiley, 2011.',
 '[10] 李华, 张伟. 基于Spring Boot与Vue的高校教务系统设计与实现[J]. 计算机时代, 2021(8): 45-49.',
 '[11] GitHub. The state of the Octoverse[R/OL]. (2023-12-31)[2026-06-01]. https://octoverse.github.com.',
 '[12] 于佳杰, 等. 面向专业认证的毕业要求达成度统一计算平台详细设计说明书[Z]. 2026.',
]
for t in refs: add_ref(doc, t)

# ===== 基金项目 / 作者简介 =====
add_para(doc, '基金项目：××××年度××教学改革研究项目"××××"（项目编号：××××）。',
         cn='宋体', size=9, indent=False, line=1.4, space_before=4)
add_para(doc, '作者简介：于佳杰，男，本科生，研究方向为软件工程与人工智能辅助开发，email：××××@××.com。',
         cn='宋体', size=9, indent=False, line=1.4)

# ===================================================================
# 英文标题 / 作者 / 摘要 / 关键词
# ===================================================================
doc.add_paragraph().paragraph_format.space_after = Pt(4)
ep = doc.add_paragraph(); ep.alignment = C.CENTER
r = ep.add_run('Exploration and Practice of a Software Development Model Based on '
               'Open-Source Collaboration and AI Assistance: A Case Study of the Unified '
               'Calculation Platform for Graduation Requirement Attainment')
set_run(r, 'Times New Roman', 'Times New Roman', 14, bold=True)

ep2 = doc.add_paragraph(); ep2.alignment = C.CENTER
r = ep2.add_run('Jiajie Yu')
set_run(r, 'Times New Roman', 'Times New Roman', 9)
ep3 = doc.add_paragraph(); ep3.alignment = C.CENTER
r = ep3.add_run('(School of Computer Science and Technology, ×× University, ×× City, ××××××, China)')
set_run(r, 'Times New Roman', 'Times New Roman', 9)

abstract_en = ('To address the long-standing reliance on manual calculation, '
 'low efficiency, error-proneness and poor cross-major reusability of graduation requirement '
 'attainment evaluation under engineering education accreditation, this paper explores a new '
 'software development model for university projects by combining open-source collaboration '
 'with AI-assisted development. Taking the "Unified Calculation Platform for Graduation '
 'Requirement Attainment" as a case study, it elaborates an open-source collaboration workflow '
 'based on Git branching strategies, pull request review and a continuous-integration toolchain, '
 'as well as a human-machine collaborative practice driven by large language models in requirement '
 'analysis, code generation, test verification and documentation. The platform adopts a '
 'front-end and back-end separated architecture with Spring Boot and Vue 3, builds a four-role '
 'RBAC model (administrator, academic administrator, major leader and lecturer), and realizes '
 'automatic three-level attainment calculation. Practice shows that the model significantly '
 'improves development efficiency and code quality, providing a reusable paradigm for software '
 'engineering teaching and the agile construction of complex information systems in universities.')
ap = doc.add_paragraph(); ap.paragraph_format.line_spacing = 1.5; first_line_indent_chars(ap, 2)
r = ap.add_run('Abstract: '); set_run(r, 'Times New Roman', 'Times New Roman', 9, bold=True)
r = ap.add_run(abstract_en); set_run(r, 'Times New Roman', 'Times New Roman', 9)

kp = doc.add_paragraph(); kp.paragraph_format.line_spacing = 1.5; first_line_indent_chars(kp, 2)
r = kp.add_run('Key words: '); set_run(r, 'Times New Roman', 'Times New Roman', 9, bold=True)
r = kp.add_run('engineering education accreditation; open-source collaboration; '
               'AI-assisted development; attainment calculation; software engineering practice')
set_run(r, 'Times New Roman', 'Times New Roman', 9)

# ----------------- 保存 -----------------
out = r'C:\Users\YU\Desktop\Computing_Platform_frontend\Computing_Platform\基于开源协作和AI辅助的软件开发模式探索与实践.docx'
doc.save(out)

# 统计中文字数（正文部分）
all_text = []
for p in doc.paragraphs:
    all_text.append(p.text)
for tb in doc.tables:
    for row in tb.rows:
        for cell in row.cells:
            all_text.append(cell.text)
full = '\n'.join(all_text)
zh = re.findall(r'[一-鿿]', full)
nonspace = len(re.sub(r'\s', '', full))
print('SAVED:', out)
print('Chinese chars:', len(zh))
print('Total chars (no whitespace ≈ Word 字符数不计空格):', nonspace)
print('Paragraphs:', len(doc.paragraphs), 'Tables:', len(doc.tables))
